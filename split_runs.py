import copy
import re
from docx import Document
from docx.text.run import Run
from docx.text.hyperlink import Hyperlink
from docx.oxml import parse_xml
from docx.enum.text import WD_BREAK
from tqdm import tqdm

def copy_styles(new_run, run):
    if not '<w:hyperlink' in run._element.xml:
        rPr_target = new_run._r.get_or_add_rPr()
        rPr_target.addnext(copy.deepcopy(run._r.get_or_add_rPr()))
        new_run._r.remove(rPr_target)
    else:
        run_copy_xml = parse_xml(run._element.xml)
        hyperlink = Hyperlink(run_copy_xml, run._parent)
        return hyperlink._element

def castom_deepcopy(runs):
    copied_runs = []
    for run in runs:
        if '<w:hyperlink' in run._element.xml:
            run_copy_xml = parse_xml(run._element.xml)
            new_run = Hyperlink(run_copy_xml, run._parent)
        else:
            run_copy_xml = parse_xml(run._element.xml)
            new_run = Run(run_copy_xml, run._parent)
        copy_styles(new_run, run)
        copied_runs.append(new_run)
    return copied_runs

def split_words_into_runs(doc_path, output_path=None):
    doc = Document(doc_path)
    if not output_path:
        output_path = doc_path
    for paragraph in tqdm(doc.paragraphs, desc='paragraphs reworked'):
        hyperlinks = paragraph.hyperlinks
        old_runs = castom_deepcopy(paragraph.iter_inner_content())

        paragraph.clear()

        full_text = ''.join(run.text for run in old_runs)
        parts = re.findall(r'\S+|\s+', full_text)

        run_index = 0
        char_index = 0
        footnote_index = 1
        for part in parts:
            while run_index < len(old_runs) and char_index >= len(old_runs[run_index].text):
                char_index -= len(old_runs[run_index].text)
                run_index += 1
            if run_index < len(old_runs):
                if '<w:hyperlink' in old_runs[run_index]._element.xml:
                    element = copy_styles(old_runs[run_index], old_runs[run_index])
                    if hyperlinks:
                        if element.rId == hyperlinks[0]._element.rId:
                            paragraph._element.append(element)
                            del hyperlinks[0]
                else:
                    new_run = paragraph.add_run(part)
                    copy_styles(new_run, old_runs[run_index])
            char_index += len(part)

        for index in range(0, len(old_runs)):
            if '<w:footnoteReference w:id="' in old_runs[index]._element.xml:
                text_til_footnote = ''.join(run.text for run in old_runs[0:index])
                parts_til_footnote = re.findall(r'\S+|\s+', text_til_footnote)
                paragraph._p.insert(len(parts_til_footnote) + footnote_index, old_runs[index]._r)
                footnote_index += 1
            if '<w:drawing>' in old_runs[index]._element.xml or '<w:pict>' in old_runs[index]._element.xml:
                paragraph._element.append(old_runs[index]._element)
            if '<w:br w:type="page"/>' in old_runs[index]._element.xml:
                paragraph.add_run().add_break(WD_BREAK.PAGE)
    doc.save(output_path)

split_words_into_runs('in.docx', 'out.docx')
